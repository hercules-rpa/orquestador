import time
import os
from datetime import datetime
import json
import requests
from docx.shared import RGBColor
from model.process.Proceso2.Entities.Congreso import Congreso
from model.process.Proceso2.Entities.Comites.ComiteAntropologiaSocial import ComiteAntropologiaSocial
from model.process.Proceso2.Entities.Comites.ComiteBiblioteconomiaYDocumentacion import ComiteBiblioteconomiaYDocumentacion
from model.process.Proceso2.Entities.Comites.ComiteArquitecturaIngenieriaCivilYUrbanismo import ComiteArquitecturaIngenieriaCivilYUrbanismo
from model.process.Proceso2.Entities.Comites.ComiteFilosofiaFilologiaYLinguistica import ComiteFilosofiaFilologiaYLinguistica
from model.process.Proceso2.Entities.Comites.ComiteCienciasEducacion import ComiteCienciasEducacion
from model.process.Proceso2.Entities.Comites.ComiteCienciasComportamiento import ComiteCienciasComportamiento
from model.process.Proceso2.Entities.Comites.ComiteCienciasSociales import ComiteCienciasSociales
from model.process.Proceso2.Entities.Comites.ComiteEstudiosGenero import ComiteEstudiosGenero
from model.process.Proceso2.Entities.Comites.ComiteCienciasEconomicasYEmpresariales import ComiteCienciasEconomicasYEmpresariales
from model.process.Proceso2.Entities.Comites.ComiteIngenierias import ComiteIngenierias
from model.process.Proceso2.Entities.Comites.ComiteQuimica import ComiteQuimica
from model.process.Proceso2.Entities.Comites.ComiteBiologiaCelularYMolecular import ComiteBiologiaCelularYMolecular
from model.process.Proceso2.Entities.Acreditacion import Acreditacion
from model.process.Proceso2.Entities.Comites.ComiteCienciasBiomedicas import ComiteCienciasBiomedicas
from model.process.Proceso2.Entities.Comites.Comite import Comite
from model.process.Proceso2.Entities.Comites.ComiteCienciasNaturaleza import ComiteCienciasNaturaleza
from model.process.Proceso2.Entities.Evaluaciones.EvaluacionSexenios.EvaluacionSexenio import EvaluacionSexenio
from model.EDMA import EDMA
from model.process.Proceso2.Entities.Evaluador import Evaluador
from model.process.Proceso2.Entities.DatosInvestigador import DatosInvestigador
from model.process.Proceso2.Entities.Enumerados.TipoIdentificadorInvestigador import TipoIdentificadorInvestigador
from model.process.Proceso2.Entities.Informe import Informe
from model.process.Proceso2.Entities.RO import RO
from model.process.ProcessCommand import ProcessCommand, ProcessID
from model.process.ProcessCommand import Pstatus as pstatus


VERSION = "1.0"
NAME = "Sexenios"
DESCRIPTION = "Proceso que genera un informe de ayuda al investigador para preparar un sexenio"
REQUIREMENTS = ["python-docx", "pandas", "sparqlwrapper", "requests"]
ID = ProcessID.SEXENIOS.value


class ProcessSexenios(ProcessCommand):
    EDMA = EDMA()

    def __init__(self, id_schedule, id_log, id_robot, priority, log_file_path, parameters=None):
        ProcessCommand.__init__(self, ID, NAME, REQUIREMENTS, DESCRIPTION,
                                id_schedule, id_log, id_robot, priority, log_file_path, parameters)


    def execute(self):
        self.state = pstatus.RUNNING
        self.log.start_log(time.time())
        self.log.state = "OK"
        self.notificar_actualizacion('Comienza la ejecuci??n del proceso.')

        if not self.parameters:
            self.notificar_actualizacion("No se han establecido par??metros")
            self.log.state = "ERROR"
            self.log.end_log(time.time())
            self.state = pstatus.FINISHED
            return

        self.notificar_actualizacion('Obteniendo el par??metro "comit??".')
        comite: Comite = self.get_comite(self.parameters['comite'])
        if not comite:
            self.notificar_actualizacion(
                'ERROR no se ha podido obtener el comit??.')
            self.log.state = "ERROR"
            self.log.end_log(time.time())
            self.state = pstatus.FINISHED
            return
        else:

            self.notificar_actualizacion('Obteniendo par??metro "per??odo".')
            periodo = self.get_periodo(self.parameters['periodo'])

            # comprobar si el comit?? no usa libros o capitulos de libros
            self.notificar_actualizacion(
                'Obteniendo par??metros relacionados con el investigador.')
            infoInvestigador = self.get_param_investigador(self.parameters)

            self.notificar_actualizacion(
                'Obteniendo los datos del investigador utilizando ED.')
            datosInvestigador = self.EDMA.get_datos_investigador(infoInvestigador)
            if not datosInvestigador:
                self.notificar_actualizacion(
                'ERROR no se ha podido obtener la informaci??n del investigador.')
                self.log.state = "ERROR"
                self.log.end_log(time.time())
                self.state = pstatus.FINISHED
                return

            self.log.completed = 10
            
            self.notificar_actualizacion(
                "Se procede a recuperar producci??n cient??fica del investigador de la base de datos de H??rcules-EDMA")

            self.notificar_actualizacion(
                'Obteniendo la lista de art??culos del investigador.')
            
            lista_articulos = self.get_articulos(
                infoInvestigador, periodo, comite.autoria_orden)        
            lista_capitulos_libros = None
            lista_libros = None
            if comite.libros_caps:
                self.notificar_actualizacion(
                    'Obteniendo los cap??tulos de libros del investigador.')
                lista_capitulos_libros = self.EDMA.get_capitulos_libros(
                    infoInvestigador, periodo)
                self.notificar_actualizacion(
                    'Obteniendo los libros del investigador.')
                lista_libros = self.EDMA.get_libros(infoInvestigador, periodo)
            
            lista_congresos:list=None
            if comite.congresos:
                self.notificar_actualizacion(
                    'Obteniendo los trabajos presentados en congresos del investigador.')
                lista_congresos = self.get_trabajos_congresos(infoInvestigador, periodo)

            lista_patentes:list = None
            if comite.patentes:
                pass
            
            self.log.completed = 40

            self.notificar_actualizacion(
                "Se proceder?? a realizar el informe para preparaci??n de sexenios de " + datosInvestigador.nombreCompleto)
            informe = Informe()
            informe.insertar_titulo('Informe de autoevaluaci??n de la investigaci??n para la preparaci??n de sexenios')
            document = informe.devolver_documento()

            p = document.add_paragraph('\n')
           
            self.notificar_actualizacion(
                'Insertando la informaci??n del investigador en el informe.')

            self.print_datos_investigador(
                p, datosInvestigador, infoInvestigador, comite, periodo)

            self.notificar_actualizacion('Obteniendo la baremaci??n en base a los requisitos del comit?? ' + comite.nombre)
            evaluacion:EvaluacionSexenio = comite.get_evaluacion_sexenio(lista_articulos)
            self.notificar_actualizacion('Insertando la baremaci??n obtenida en el informe.')
            self.print_baremacion(document,evaluacion)
            
            principales = []
            sustitutorios = []
            if lista_articulos:
                if evaluacion.produccion_principal:
                    principales = evaluacion.produccion_principal
                elif len(lista_articulos) > 5:
                    principales = lista_articulos[0:5]
                else:
                    principales = lista_articulos
                
                sustitutorios = evaluacion.get_produccion_sustitutoria()
                if len(sustitutorios) > 30:
                    sustitutorios = sustitutorios[0:30]                

            self.log.completed = 50

            self.notificar_actualizacion(
                "Insertando en el informe la producci??n cient??fica principal.")

            self.print_produccion_cientifica(
                document, principales, comite, 'principal')
            self.log.completed = 60

            self.notificar_actualizacion(
                "Insertando en el informe la producci??n cient??fica sustitutoria.")
            self.log.completed = 70
            self.print_produccion_cientifica(
                document, sustitutorios, comite, 'sustitutoria')

            if comite.alerta_autores:
                # LIBROS
                self.notificar_actualizacion(
                    'Insertando en el informe los libros.')
                self.print_libros(document, lista_libros, comite)

                # CAPITULOS DE LIBROS
                self.notificar_actualizacion(
                    'Insertando en el informe los cap??tulos de libros.')
                self.print_capitulos_libros(
                    document, lista_capitulos_libros, comite)

            self.log.completed = 80

            if comite.congresos:
                self.notificar_actualizacion(
                    'Insertando en el informe los trabajos presentados en congresos.')
                self.print_trabajos_congresos(document, lista_congresos)

            if comite.patentes:
                pass

            self.log.completed = 90

            document.add_page_break()

            self.notificar_actualizacion("Persistiendo informe...")
            doc_name = self.save_informe(
                document, datosInvestigador.nombreCompleto, True)
            self.notificar_actualizacion(
                "Informe generado y guardado con el nombre "+doc_name + ".")
            self.log.completed = 90

            self.notificar_actualizacion("Subiendo informe al CDN...")
            response = None

            #response = self.upload_file(doc_name)
            if response and  response.status_code == 201:
                json_dicti = json.loads(response.text)
                self.notificar_actualizacion(
                    "Fichero subido correctamente, url de descarga: "+json_dicti["URL"])
            else:
                self.notificar_actualizacion('ERROR al subir el informe al CDN.')

        self.notificar_actualizacion(
            'Proceso de generaci??n de informe de sexenio finalizado.')
        self.log.completed = 100
        self.log.end_log(time.time())
        self.state = pstatus.FINISHED

    def pause(self):
        pass

    def kill(self):
        pass

    def resume(self):
        pass

    def get_autores_formatted(self, autoresPandas):
        """
        M??todo que formatea el nombre de los autores.
        """
        autores = ''
        for name in autoresPandas.loc[:, "nombreAutor.value"]:
            autores += name+', '
        return autores[:-2]
    
    def get_comite(self, id:str) -> Comite:
        """
        M??todo que crea el comit?? que eval??a la solicitud del sexenio.
        """
        result:Comite = None
        currentdir = os.path.dirname(os.path.realpath(__file__))
        conf = self.get_configuraciones(currentdir+'/Configuraciones/comite.json')
        if id == "2":
            result = ComiteQuimica(id, conf[id])        
        elif id == "3":
            result = ComiteBiologiaCelularYMolecular(id, conf[id])
        elif id == "4":
            result = ComiteCienciasBiomedicas(id, conf[id])
        elif id == "5":
            result = ComiteCienciasNaturaleza(id=id, evaluador=conf[id])
        elif id == "7":
            result = ComiteIngenierias(id=id, evaluador=conf[id])
        elif id == "8":
            tecnologico = False
            if 'perfil_tecnologico' in self.parameters:
                tecnologico = self.parameters['perfil_tecnologico']
            
            self.notificar_actualizacion('Perfil tecnol??gico: ' + str(tecnologico))
            result = ComiteArquitecturaIngenieriaCivilYUrbanismo(id, conf[id], False, tecnologico) 
        
        elif id == "9":
            if 'subcomite' in self.parameters:
                subcomite = self.parameters['subcomite']             
                if subcomite == 1:
                    result = ComiteCienciasSociales(id, conf[id])
                elif subcomite == 2 or subcomite == 6:  
                    #El subcomit?? OTROS tiene los mismos requisitos.
                    result = ComiteCienciasComportamiento(id, conf[id])
                elif subcomite == 3:
                    result = ComiteBiblioteconomiaYDocumentacion(id, conf[id])
                elif subcomite == 4:
                    result = ComiteEstudiosGenero(id, conf[id])
                elif subcomite == 5:
                    result = ComiteAntropologiaSocial(id, conf[id])
        elif id == "10":
            result = ComiteCienciasEducacion(id, conf[id])     
        elif id == "11":
            result = ComiteCienciasEconomicasYEmpresariales(id, conf[id])
        elif id == "14":
            result = ComiteFilosofiaFilologiaYLinguistica(id, conf[id])  
        return result

    def get_configuraciones(self, filename):
        """
        M??todo que obtiene la configuraci??n del robot.
        """
        dirname = os.path.dirname(__file__)
        file_path = os.path.join(dirname, filename)
        try:
            with open(file_path, encoding="utf-8") as json_robot:
                config = json.load(json_robot)
                return config
        except:
            print("Error No such file or directory")
        return None

    def get_param_investigador(self, parameters):
        """
        M??todo que obtiene el tipo de identificador del investigador y el identificador 
        de los par??metros que recibe el proceso para su ejecuci??n.
        """
        investigador: str = ""
        tipoId: TipoIdentificadorInvestigador = TipoIdentificadorInvestigador.NODEFINIDO
        try:
            if 'tipoId' in parameters:
                tipoId = TipoIdentificadorInvestigador(
                    parameters['tipoId'])

            if 'investigador' in parameters:
                investigador = parameters['investigador']

        except Exception as e:
            print(e)
            self.notificar_actualizacion(
                'ERROR al obtener el identificador del investigador.')

        return (tipoId, investigador)

    def get_periodo_array(self, period):
        """
        M??todo que obtiene calcula el per??odo de a??os en base a un string. 
        Ejemplo: '1,2,5-7' a [1,2,5,6,7]
        """
        result = []
        try:
            for part in period.split(','):
                if '-' in part:
                    a, b = part.split('-')
                    a, b = int(a), int(b)
                    result.extend(range(a, b + 1))
                else:
                    a = int(part)
                    result.append(a)

        except Exception as e:
            print(e)
            self.notificar_actualizacion(
                'ERROR al obtener el par??metro "periodo".')
        return result

    def get_periodo(self, periodo):
        """
        M??todo que obtiene los a??os que forman parte del per??odo separados por comas.
        """
        result = self.get_periodo_array(periodo)
        periodo = ''
        if result:
            for a in result:
                periodo += str(a)+','

            periodo = periodo[:-1]
            print('Periodo: ' + periodo)
        return periodo

    def get_articulo(self, id_articulo) -> RO:
        """
        M??todo que obtiene un art??culo e hidrata la entidad RO.
        """
        articulo: RO = None
        try:
            articulo = self.EDMA.get_articulo(id_articulo)
            if articulo:
                autores = self.EDMA.get_autores_articulo(id_articulo)
                articulo.autores = self.get_autores_formatted(autores)
                articulo.nautores = len(autores.loc[:, "nombreAutor.value"])
        except:
            self.notificar_actualizacion(
                'ERROR a obtener el art??culo con id: ' + id_articulo)

        return articulo

    def get_trabajo_congreso(self, id) -> Congreso:
        """
        M??todo que obtiene un trabajo presentado en un congreso e hidrata la entidad Congreso.
        """
        congreso: Congreso = None
        try:
            congreso = self.EDMA.get_trabajo_congreso(id)
            if congreso:
                autores = self.EDMA.get_autores_trabajo_congreso(id)
                congreso.autores = self.get_autores_formatted(autores)
                congreso.num_autores = len(autores.loc[:, "nombreAutor.value"])
        except:
            self.notificar_actualizacion(
                'ERROR a obtener el trabajo presentado en un congreso con id: ' + id)

        return congreso

    def get_articulos(self, infoInvestigador, periodo, autoria_orden, max_article:int=0) -> list:
        """
        M??todo que obtiene la lista de art??culos de un investigador.
        """
        result = []
        dataframe_art = self.EDMA.get_lista_articulos_investigador(
            infoInvestigador, max_article, periodo, autoria_orden)
        if not dataframe_art.empty:
            self.notificar_actualizacion('Obteniendo la informaci??n individual de la lista de art??culos.')
            for i in range(len(dataframe_art)):
               articulo = self.get_articulo(dataframe_art.iloc[i]['doc.value'])
               if articulo: 
                    articulo.posicion_autor = dataframe_art.iloc[i]['posicion.value']
                    result.append(articulo)

            print('Se han obtenido ' + str(len(result)) + ' art??culos.')            
        return result

    def get_trabajos_congresos(self, infoInvestigador, periodo) -> list:
        """
        M??todo que obtiene la lista de trabajos presentados en congresos de un investigador.
        """
        result = []
        dataframe_art = self.EDMA.get_trabajos_congresos(infoInvestigador, periodo)
        if not dataframe_art.empty:
            self.notificar_actualizacion('Obteniendo la informaci??n individual de la lista de trabajos presentados en congresos.')
            for i in range(len(dataframe_art)):
               congreso = self.get_trabajo_congreso(dataframe_art.iloc[i]['doc.value'])
               if congreso: 
                    result.append(congreso)            
        return result

    def print_datos_investigador(self, p, datosInvestigador: DatosInvestigador, 
        infoInvestigador, evaluador:Evaluador, periodo: str = '', acreditacion:Acreditacion=None):
        """
        M??todo que imprime los datos del investigador.
        """
        p.add_run('Nombre y apellidos del Investigador: ' +
                  datosInvestigador.nombreCompleto+'\n')
        p.add_run('Universidad: '+datosInvestigador.nombreUniversidad+'\n')
        p.add_run('Departamento: '+datosInvestigador.nombreDepartamento+'\n')
        p.add_run('Email: '+datosInvestigador.email+' \n')
        if infoInvestigador[0] == TipoIdentificadorInvestigador.ORCID:
            p.add_run('ORCID: '+infoInvestigador[1]+'. \n')
        if periodo:
            p.add_run('Per??odo del sexenio: '+periodo+'.\n')
        
        if evaluador and evaluador.comision:
            p.add_run('Comisi??n: '+evaluador.nombre+'. \n')
        else:
            p.add_run('Comit??: '+evaluador.nombre+'. \n')

        if evaluador.autoria_str:
            p.add_run('Relevancia Autor??as: '+evaluador.autoria_str+'. \n')
        
        if acreditacion and acreditacion.nombre:
            p.add_run('Tipo de Acreditaci??n: '+acreditacion.nombre+'. \n')

    def print_baremacion(self, documento, evaluacion:EvaluacionSexenio):
        """
        M??todo que imprime la baremaci??n obtenida tras evaluar la lista de art??culos.
        """
        documento.add_heading('Baremaci??n m??nima obtenida:')
        
        p = documento.add_paragraph('En la baremaci??n que se ha realizado solo se ha tenido en cuenta los requisitos m??nimos obligatorios ' + \
                'de cada comit??. \n\n')

        if evaluacion:
            if evaluacion.puntuacion > 0:
                p.add_run('Se ha realizado una baremaci??n de la producci??n cient??fica principal ' + \
                    'en base a los criterios de la ANECA. \n' + \
                    'La puntuaci??n m??nima que podr??a obtener seg??n los requisitos es: ' + \
                    str(evaluacion.puntuacion) + '. \n')
            else:
                p.add_run('No se ha podido obtener una baremaci??n con las puntuaciones de la producci??n cient??fica principal. \n')

            if evaluacion.observaciones:
                p.add_run(evaluacion.observaciones)

                
    def print_articulo(self, document, articulo, roposition, evaluador:Evaluador) -> bool:
        """
        M??todo encargado de consultar la informaci??n de un art??culo e insertarlo en el informe, 
        retorna true si se ha incluido correctamente el art??culo o False si no ha sido posible.
        """
        try:
            if articulo:
                document.add_heading('Posici??n: '+str(roposition), level=1)
                pais_edicion = 'Espa??a'
                npaginas = str(articulo.get_npaginas())
                
                document.add_heading('Informaci??n general del art??culo:', level=3)
                table = document.add_table(rows=0, cols=1)
                table.style = 'Table Grid'
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('Revista: ').bold = True
                d.add_run(articulo.revista+'\n')
                d.add_run('Editorial: ').bold = True
                d.add_run(articulo.editorial+'\n')
                d.add_run('Pa??s de Edici??n: ').bold = True
                d.add_run(pais_edicion)
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('T??tulo: ').bold = True
                d.add_run(articulo.titulo+'\n')
                d.add_run('Autores: ').bold = True
                d.add_run(str(articulo.nautores))
                
                if(evaluador.alerta_autores != -1 and articulo.nautores > evaluador.alerta_autores):
                    d.add_run('\n')
                    eval = 'comit??'                    
                    if evaluador.comision:
                        eval = 'comisi??n'

                    run = d.add_run('N??mero de autores ('+str(articulo.nautores)+') elevado por encima de lo recomendado para este ' + eval + ' ('+str(evaluador.alerta_autores) +
                                    ') , en este ' + eval + ' se tiene muy en cuenta el n??mero de autores, justificar muy bien la aportaci??n y carga de trabajo en el art??culo.')
                    
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('A??o: ').bold = True
                d.add_run(articulo.get_anio()+'\n')
                d.add_run('Volumen: ').bold = True
                d.add_run(articulo.volumen+'\n')
                d.add_run('N??mero: ').bold = True
                d.add_run(str(articulo.numero)+'\n')
                d.add_run('N??mero de p??ginas: ').bold = True
                d.add_run(npaginas+'\n')
                d.add_run('P??gina inicio: ').bold = True
                d.add_run(articulo.pag_inicio+'\n')
                d.add_run('P??gina fin: ').bold = True
                d.add_run(articulo.pag_fin+'\n')

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('ISSN: ').bold = True
                d.add_run(articulo.issn)

                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('DOI: ').bold = True
                d.add_run(articulo.doi)

                document.add_heading('\nInformaci??n JCR del Art??culo', level=3)
                # Puede estar en varias categor??as, poner las dos en que cuartil est?? o poner mejor la m??s cercana a nuestro ??mbito cient??fico
                ctematica = "No definida"
                table = document.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells

                d = hdr_cells[0].add_paragraph()
                d.add_run('JCR (Journal Citation Reports)\n').bold = True
                d.add_run('??ndice de impacto de SCIE y SSCIE de Web Of Science Colecci??n Principal. \n ' +
                          'Revistas multidisciplinares y cobertura mundial\nRecurso indicado en CNEAI 2021 en los Campos 1-8, 10 y 11.\n\n')
                d.add_run('Tiene un factor de impacto de ')
                d.add_run(str(articulo.impacto)).bold = True
                d.add_run(' en JCR, A??o ')
                d.add_run(articulo.get_anio()+'\n').bold = True
                d.add_run('Ocupa la posici??n ')
                d.add_run(str(articulo.posicion)).bold = True
                d.add_run(' de un total de ')
                d.add_run(str(articulo.n_revistas)).bold = True
                d.add_run(' revistas en la categor??a tem??tica ')
                d.add_run(ctematica).bold = True
                d.add_run(' Por lo que est?? en el cuartil ')
                d.add_run(articulo.cuartil + '??.').bold = True

                document.add_heading(
                    '\nWeb of Science Colecci??n Principal (WOS CC)', level=3)

                d = document.add_paragraph()
                if articulo.wos_cites:
                    d.add_run('Total N?? de citas: ' +
                              articulo.wos_cites+'\n').bold = True
                else:
                    d.add_run('No se encontraron citas en WOS\n').bold = True

                document.add_heading('\nSCOPUS', level=3)
                d = document.add_paragraph()

                if articulo.scopus_cites:
                    d.add_run('Total N?? de citas: ' +
                              articulo.scopus_cites+'\n').bold = True
                else:
                    d.add_run('No se encontraron citas en SCOPUS\n').bold = True

                document.add_heading('\nSemantic Scholar', level=3)
                d = document.add_paragraph()

                if articulo.ss_cites:
                    d.add_run('Total N?? de citas: ' +
                              articulo.ss_cites+'\n').bold = True
                else:
                    d.add_run(
                        'No se encontraron citas en Semantic Scholar\n').bold = True
                document.add_paragraph('\n')

            return True
        except Exception as e:
            print(e)
            self.notificar_actualizacion(
                'ERROR en la inserci??n de un art??culo en el informe.')
        return False

    def print_produccion_cientifica(self, document, articulos:list, evaluador, tipoProduccion: str):
        """
        M??todo que imprime la produccion cientifica de un investigador.
        """
        if articulos:
            roposition = 1

            if tipoProduccion == 'principal':
                document.add_heading(
                    'Ranking de Producci??n Cient??fica : Producci??n principal.', level=1)
            else:
                document.add_paragraph(
                    'A continuaci??n se detallan hasta 30 art??culos sustitutorios por si el investigador desea sustituir producci??n cient??fica calificada como principal:')
            for articulo in articulos:
                if self.print_articulo(document, articulo, roposition, evaluador):
                    roposition += 1
        else:
            document.add_heading(
                'No se encontr?? producci??n cient??fica ' + tipoProduccion + '.', level=3)
            self.notificar_actualizacion(
                'No se encontr?? producci??n cient??fica ' + tipoProduccion + '.')

    def print_libros(self, document, libros, evaluador):
        """
        M??todo encargado de insertar los libros en el informe.
        """
        document.add_heading(
            'Apartado de libros y monograf??as cient??ficas', level=1)

        if not libros.empty:
            eval = 'el comit?? seleccionado'
            if evaluador.comision:
                eval = 'la comisi??n seleccionada'
            document.add_paragraph(
                'A continuaci??n se detallan libros o monograf??as cient??ficas que el autor ha llevado a cabo, ya que en ' + eval + ': '+evaluador.nombre+' pueden ser valorados.')
            document.add_paragraph()
            for i in range(len(libros)):
                document.add_paragraph(
                    '-T??tulo: '+libros.iloc[i]['titulo.value']+'. '+'Fecha de publicaci??n: '+libros.iloc[i]['anioFecha.value'], style='List Bullet')
        else:
            document.add_heading(
                'No se encontraron libros', level=3)
            self.notificar_actualizacion(
                'No se encontraron libros')

    def print_capitulos_libros(self, document, capitulos, evaluador):
        """
        M??todo que inserta los cap??tulos de libros en el informe.
        """
        document.add_heading('Apartado de cap??tulos de libros', level=1)

        if not capitulos.empty:
            eval = 'el comit?? seleccionado'
            if evaluador.comision:
                eval = 'la comisi??n seleccionada'
            document.add_paragraph(
                'A continuaci??n se detallan los cap??tulos de libros que el autor ha llevado a cabo, ya que en ' + eval + ': '+evaluador.nombre+' pueden ser valorados.')
            
            d = document.add_paragraph()
            for i in range(len(capitulos)):
                document.add_paragraph('-T??tulo: '+capitulos.iloc[i]['titulo.value']+'. ' +
                                       'Fecha de publicaci??n: '+capitulos.iloc[i]['anioFecha.value'], style='List Bullet')

        else:
            document.add_heading(
                'No se encontraron cap??tulos de libros', level=3)
            self.notificar_actualizacion(
                'No se encontratron cap??tulos de libros.')

    def print_trabajos_congresos(self, document, elements):
        """
        M??todo encargado de insertar los trabajos presentados en congresos 
        nacionales e internacionales en el informe.
        """
        document.add_heading(
            'Apartado de trabajos presentados en congresos nacionales e internacionales.', level=1)

        if elements:
            posicion = 1
            for element in elements:
                document.add_heading('Posici??n: '+str(posicion), level=1)
                document.add_heading('Informaci??n general del trabajo:', level=3)
                table = document.add_table(rows=0, cols=1)
                table.style = 'Table Grid'
                row_cells = table.add_row().cells
                d = row_cells[0].add_paragraph()
                d.add_run('T??tulo: ').bold = True
                d.add_run(element.titulo+'\n')
                if element.autores:
                    d.add_run('Autores: ').bold = True
                    d.add_run(element.autores+'\n')
                    d.add_run('N??mero de autores: ').bold = True
                    d.add_run(str(element.num_autores)+'\n')
                if element.fecha:
                    d.add_run('Fecha publicaci??n: ').bold = True
                    d.add_run(element.fecha+'\n')

                posicion += 1

        else:
            document.add_heading(
                'No se encontraron trabajos presentados en congresos.', level=3)
            self.notificar_actualizacion('No se encontraron trabajos presentados en congresos.')

    def save_informe(self, document, nombreInvestigador, is_sexenio:bool) -> str:
        """
        M??todo que calcula el nombre del informe y lo almacena en la ubicaci??n correspondiente. 
        Retorna el nombre del documento guardado.
        """
        date_time = datetime.now().strftime("%m-%d-%Y-%H%M%S")
        doc_name = 'Informe'
        if is_sexenio:
            doc_name +='Sexenios'
        else:
            doc_name+='Acreditacion'

        doc_name += nombreInvestigador+date_time+'.docx'
        doc_name = doc_name.replace(' ', '')
        document.save(doc_name)
        return doc_name

    def upload_file(self, filename):
        url = "http://10.208.99.12:5000/api/orchestrator/files"

        payload = {}
        files = [
            ('file', (filename, open(filename, 'rb'), 'application/docx'))
        ]
        headers = {}

        response = requests.request(
            "POST", url, headers=headers, data=payload, files=files)
        return response
